# -*- coding: utf-8 -*-
"""End-to-end тест ядра: детекция, замена, docx, восстановление."""

import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from anon import engine, readers, writers
from anon.detectors import detect_orgs_regex, detect_structured, valid_inn, valid_ogrn, valid_snils
from anon.entities import Entity, TagRegistry

SAMPLE = """АРБИТРАЖНЫЙ СУД ГОРОДА МОСКВЫ
Дело № А40-12345/2024

ИСКОВОЕ ЗАЯВЛЕНИЕ

Истец: Иванов Иван Петрович, паспорт серия 45 04 123456,
СНИЛС 112-233-445 95, проживающий по адресу: г. Москва, ул. Ленина, д. 5, кв. 12,
тел. +7 (916) 123-45-67, email: ivanov.ip@example.com.

Ответчик: ООО «Ромашка», ИНН 7707083893, ОГРН 1027700132195,
КПП 773601001, р/с 40702810900000005555, БИК 044525225.

Иванов И.П. заключил с ООО «Ромашка» договор поставки от 15 марта 2024 г.
Ответчик обязательства не исполнил до 15.03.2024, в связи с чем Иванову Ивану Петровичу
причинены убытки. Представитель истца — Сидорова Анна Викторовна (ИНН 500100732259).

Прошу взыскать с ООО «Ромашка» в пользу Иванова И.П. 1 000 000 рублей.
"""


def banner(s):
    print("\n" + "=" * 20, s, "=" * 20)


def test_helpers():
    assert valid_inn("7707083893") and valid_inn("500100732259")
    assert valid_ogrn("1027700132195")
    assert not valid_ogrn("1187746123456")

    ogrn_ctx = detect_structured("ОГРН 1187746123456")
    assert any(e.type == "OGRN" and e.text == "1187746123456" for e in ogrn_ctx)

    assert valid_snils("112-233-445 95")

    kpp = detect_structured("Организация, КПП 7701AB001, далее по тексту.")
    assert any(e.type == "KPP" and "7701AB001" in e.text for e in kpp)

    dates = detect_structured(
        "Договор от 15 марта 2024 г., срок до 15.03.2024 и 2024-03-15."
    )
    date_keys = {e.norm_key for e in dates if e.type == "DATE"}
    assert date_keys == {"2024-03-15"}
    assert all(e.tag == "" for e in dates)

    city_addrs = detect_structured(
        "Арбитражный суд города Москвы рассмотрел дело. "
        "Адрес филиала: г. Москвы. Второй адрес: г Москва. "
        "Третий адрес указан в городе Москве. Четвертый: гор. Москвы."
    )
    addr_texts = {e.text for e in city_addrs if e.type == "ADDR"}
    assert "города Москвы" in addr_texts
    assert "г. Москвы" in addr_texts
    assert "г Москва" in addr_texts
    assert "городе Москве" in addr_texts
    assert "гор. Москвы" in addr_texts

    prospect = detect_structured("проживает по адресу: Невский пр-т, д. 28, кв. 7.")
    assert any(
        e.type == "ADDR" and "Невский пр-т" in e.text and "д. 28" in e.text
        for e in prospect
    )

    by_addr = detect_structured(
        "проживает по адресу:\n220030\nпроспект Независимости,\n32А-1."
    )
    by_texts = {e.text for e in by_addr if e.type == "ADDR"}
    assert any("220030" in t for t in by_texts)
    assert any("проспект Независимости" in t and "32А-1" in t for t in by_texts)

    sites = detect_structured(
        "Сайт банка: www.sber-bank.by, зеркало https://sber-bank.by/credits "
        "и почта ivanov.ip@example.com."
    )
    site_texts = {e.text for e in sites if e.type == "SITE"}
    assert "www.sber-bank.by" in site_texts
    assert any(t.startswith("https://sber-bank.by") for t in site_texts)
    assert all("@" not in t for t in site_texts)
    assert not any(e.type == "SITE" and e.text == "bank.by" for e in sites)
    site_keys = {e.norm_key for e in sites if e.type == "SITE"}
    assert "sber-bank.by" in site_keys

    zhlobin = "г. Рогачеве, г. Жлобине ОАО «БПС-Сбербанк»"
    zhlobin_ents = engine.resolve_overlaps(
        detect_structured(zhlobin) + detect_orgs_regex(zhlobin)
    )
    assert any(e.type == "ORG" and "Сбербанк" in e.text for e in zhlobin_ents)
    addr_join = " ".join(e.text for e in zhlobin_ents if e.type == "ADDR")
    assert "Рогачеве" in addr_join and "Жлобине" in addr_join
    assert not any(e.type == "ADDR" and "ОАО" in e.text for e in zhlobin_ents)

    a = [Entity("FIO", 0, 21, "Иванов Иван Петрович", "иванов|и")]
    b = [Entity("FIO", 0, 6, "Иванов", "иванов")]
    registry = TagRegistry()
    engine.apply_package_tags([a, b], registry)
    assert a[0].tag == b[0].tag == "[ФИО1]"

    mapping = {
        "tags": {
            "[ФИО1]": {"type": "ФИО", "canonical": r"ref\1value"},
            "[ФИО10]": {"type": "ФИО", "canonical": "Петров"},
        }
    }
    restored, leftover = engine.deanonymize_text("[ФИО1] и [ФИО10]", mapping)
    assert restored == r"ref\1value и Петров"
    assert not leftover
    print("Вспомогательные проверки: ok")


def test_pdf_keeps_pdf_output():
    import fitz
    from anon.readers import LoadedDoc

    text = "Истец Иванов, ИНН 7707083893."
    loaded = LoadedDoc("contract.pdf", "pdf", text)
    start = text.index("7707083893")
    ent = Entity("INN", start, start + 10, "7707083893",
                 norm_key="7707083893", tag="[ИНН1]")
    out_name = writers.anon_output_path(loaded.path)
    assert out_name.endswith("contract_anon.pdf")
    data = writers.anonymized_bytes(loaded, [ent], out_name)
    assert data.startswith(b"%PDF")
    doc = fitz.open(stream=data, filetype="pdf")
    extracted = "".join(page.get_text("text") for page in doc)
    doc.close()
    assert "7707083893" not in extracted
    assert "[ИНН1]" in extracted
    assert writers.anon_output_path("file.docx").endswith("_anon.docx")
    assert writers.anon_output_path("file.txt").endswith("_anon.docx")
    assert writers.anon_output_path("file.pdf", ext=".txt").endswith("_anon.txt")
    assert writers.restored_output_path("answer.docx").endswith("_restored.docx")
    assert writers.restored_output_path("answer.pdf", ext=".txt").endswith(
        "_restored.txt"
    )
    txt = writers.anonymized_bytes(loaded, [ent], "contract_anon.txt")
    assert txt.startswith("Истец Иванов, ИНН [ИНН1].".encode("utf-8"))
    docx_bytes = writers.anonymized_bytes(loaded, [ent], "contract_anon.docx")
    assert docx_bytes.startswith(b"PK")


def test_csv_keeps_csv_output():
    from anon.readers import LoadedDoc

    text = "ФИО,ИНН\nИванов,7707083893\n"
    loaded = readers.load_from_bytes("people.csv", text.encode("utf-8"))
    assert loaded.kind == "csv"
    assert "Иванов" in loaded.text

    start = loaded.text.index("Иванов")
    inn_start = loaded.text.index("7707083893")
    ents = [
        Entity("FIO", start, start + 6, "Иванов",
               norm_key="иванов", tag="[ФИО1]"),
        Entity("INN", inn_start, inn_start + 10, "7707083893",
               norm_key="7707083893", tag="[ИНН1]"),
    ]
    out_name = writers.anon_output_path(loaded.path)
    assert out_name.endswith("people_anon.csv")
    data = writers.anonymized_bytes(loaded, ents, out_name)
    out = data.decode("utf-8-sig")
    assert "Иванов" not in out
    assert "7707083893" not in out
    assert "[ФИО1]" in out and "[ИНН1]" in out
    assert out.splitlines()[0] == "ФИО,ИНН"

    cp1251 = "ФИО;ИНН\nПетров;7707083893\n".encode("cp1251")
    loaded_1251 = readers.load_from_bytes("people.csv", cp1251)
    assert "Петров" in loaded_1251.text

    assert writers.anon_output_path("file.csv").endswith("_anon.csv")
    assert writers.restored_output_path("answer.csv").endswith("_restored.csv")
    restored, leftover = engine.deanonymize_text(out, {
        "tags": {
            "[ФИО1]": {"type": "ФИО", "canonical": "Иванов"},
            "[ИНН1]": {"type": "ИНН", "canonical": "7707083893"},
        }
    })
    assert not leftover
    assert "Иванов" in restored and "7707083893" in restored


def main():
    test_helpers()

    analyzer = engine.Analyzer()
    entities = analyzer.detect(SAMPLE)
    banner("Найденные сущности")
    for e in entities:
        print(f"  {e.type:9} {e.text!r:55} key={e.norm_key}")

    registry = TagRegistry()
    engine.apply_tags(entities, registry)
    anon_text = engine.anonymize_text(SAMPLE, entities)
    banner("Обезличенный текст")
    print(anon_text)

    # проверки
    assert "А40-12345/2024" in anon_text, "номер дела должен сохраниться"
    assert "АРБИТРАЖНЫЙ СУД" in anon_text.upper(), "суд не обезличиваем"
    for pii in (
        "Иванов",
        "7707083893",
        "1027700132195",
        "112-233-445",
        "40702810900000005555",
        "ivanov.ip@example.com",
        "123-45-67",
        "Сидорова",
        "Ромашка",
        "45 04 123456",
        "500100732259",
        "ул. Ленина",
        "15.03.2024",
        "15 марта 2024",
        "773601001",
    ):
        assert pii not in anon_text, f"НЕ ЗАМЕНЕНО: {pii}"

    date_tags = {e.tag for e in entities if e.type == "DATE" and e.norm_key == "2024-03-15"}
    assert len(date_tags) == 1, f"одна дата получила разные теги: {date_tags}"

    # консистентность: Иванов везде один тег
    fio_tags = {e.tag for e in entities if e.type == "FIO" and "иванов" in e.norm_key}
    assert len(fio_tags) == 1, f"Иванов получил разные теги: {fio_tags}"
    print("\nПроверки замены: ok; тег Иванова:", fio_tags)

    # round-trip
    mapping = registry.to_mapping()
    restored, leftover = engine.deanonymize_text(anon_text, mapping)
    assert not leftover, f"остались теги: {leftover}"
    assert "Иванов Иван Петрович" in restored
    assert "ООО «Ромашка»" in restored
    print("Восстановление: ok")

    # docx round-trip с форматированием
    banner("docx")
    import docx

    d = docx.Document()
    p = d.add_paragraph()
    p.add_run("Истец ")
    r = p.add_run("Иванов Иван Петрович")
    r.bold = True
    p.add_run(" предъявил иск к ")
    r2 = p.add_run("ООО «Ромашка»")
    r2.italic = True
    p.add_run(", ИНН 7707083893.")

    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "СНИЛС"
    table.rows[0].cells[1].text = "112-233-445 95"

    src = os.path.join(os.path.dirname(__file__), "sample.docx")
    d.save(src)

    loaded = readers.load(src)
    ents = analyzer.detect(loaded.text)
    reg2 = TagRegistry()
    engine.apply_tags(ents, reg2)
    out = os.path.join(os.path.dirname(__file__), "sample_anon.docx")
    writers.save_anonymized(loaded, ents, out)

    check = readers.load(out)
    print(check.text)
    assert "Иванов" not in check.text and "Ромашка" not in check.text
    assert "7707083893" not in check.text and "112-233-445" not in check.text
    # форматирование сохранилось: жирный run содержит тег ФИО
    d2 = docx.Document(out)
    bold_runs = [r.text for p2 in d2.paragraphs for r in p2.runs if r.bold]
    assert any("[ФИО" in t for t in bold_runs), f"жирный тег не найден: {bold_runs}"
    print("docx: ok (форматирование сохранено)")

    # адрес, разрезанный переносом абзаца
    split = os.path.join(os.path.dirname(__file__), "sample_split.docx")
    ds = docx.Document()
    ds.add_paragraph("Адрес: г. Москва, ул. Ленина,")
    ds.add_paragraph("д. 5, кв. 12.")
    ds.save(split)
    loaded_s = readers.load(split)
    ents_s = analyzer.detect(loaded_s.text)
    assert any(e.type == "ADDR" for e in ents_s), "адрес на стыке абзацев не найден"
    engine.apply_tags(ents_s, TagRegistry())
    out_s = os.path.join(os.path.dirname(__file__), "sample_split_anon.docx")
    writers.save_anonymized(loaded_s, ents_s, out_s)
    check_s = readers.load(out_s)
    assert "Ленина" not in check_s.text
    print("docx на стыке абзацев: ok")

    banner("pdf")
    test_pdf_keeps_pdf_output()
    print("pdf: ok")

    banner("csv")
    test_csv_keeps_csv_output()
    print("csv: ok")

    print("\nВСЕ ТЕСТЫ ПРОЙДЕНЫ")


if __name__ == "__main__":
    main()

